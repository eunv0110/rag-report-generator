#!/usr/bin/env python3
"""문서 생성 모듈 (Word/PDF)

JSON 보고서 데이터를 Word/PDF 문서로 변환
"""

from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import re
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False


class DocumentGenerator:
    """Word/PDF 문서 생성기"""

    def __init__(self):
        self.first_heading_added = False

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """헤딩 추가"""
        heading = doc.add_heading(text, level=level)
        return heading

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False, italic: bool = False):
        """단락 추가"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Malgun Gothic'

        if bold:
            run.bold = True
        if italic:
            run.italic = True

        return para

    def _parse_markdown_table(self, table_text: str) -> List[List[str]]:
        """마크다운 테이블 파싱

        Args:
            table_text: 마크다운 테이블 텍스트

        Returns:
            2D 리스트 (행x열)
        """
        lines = table_text.strip().split('\n')
        rows = []

        for i, line in enumerate(lines):
            # 구분선 스킵 (예: |---|---|)
            if i == 1 and re.match(r'^\s*\|[\s\-:]+\|\s*$', line):
                continue

            # 셀 추출
            cells = [cell.strip() for cell in line.split('|')]
            # 앞뒤 빈 셀 제거
            cells = [cell for cell in cells if cell]

            if cells:
                rows.append(cells)

        return rows

    def _add_markdown_table(self, doc: Document, table_text: str):
        """마크다운 테이블을 Word 테이블로 변환

        Args:
            doc: Word 문서
            table_text: 마크다운 테이블 텍스트
        """
        rows_data = self._parse_markdown_table(table_text)

        if not rows_data:
            return

        # Word 테이블 생성
        num_rows = len(rows_data)
        num_cols = len(rows_data[0]) if rows_data else 0

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # 데이터 채우기
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_data

                    # 첫 행은 헤더로 볼드 처리
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                    else:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)

        doc.add_paragraph()  # 테이블 뒤 간격

    def _add_formatted_text(self, doc: Document, text: str):
        """마크다운 형식이 포함된 텍스트를 Word에 추가

        Args:
            doc: Word 문서
            text: 마크다운 형식 텍스트
        """
        # 테이블 감지 및 처리
        table_pattern = r'(\|.+\|(?:\n\|.+\|)+)'
        parts = re.split(table_pattern, text)

        for part in parts:
            if not part.strip():
                continue

            # 테이블인 경우
            if re.match(r'^\s*\|', part):
                self._add_markdown_table(doc, part)
            else:
                # 일반 텍스트 처리
                self._add_formatted_paragraph(doc, part)

    def _add_formatted_paragraph(self, doc: Document, text: str):
        """마크다운 형식(볼드, 이탤릭, 리스트)을 Word로 변환

        Args:
            doc: Word 문서
            text: 텍스트
        """
        lines = text.strip().split('\n')

        for line in lines:
            if not line.strip():
                continue

            # 수평선 감지 (---, ___, ***)
            if re.match(r'^[\-_*]{3,}\s*$', line):
                # 수평선은 스킵 (표시하지 않음)
                continue

            # 리스트 항목 감지
            list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                content = list_match.group(3)
                para = doc.add_paragraph(style='List Bullet')
                self._add_inline_formatting(para, content)
                continue

            # 헤딩 감지
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                content = heading_match.group(2)
                heading = doc.add_heading(content, level=level)

                # 첫 번째 레벨 1 헤딩만 중앙 정렬
                if not self.first_heading_added and level == 1:
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self.first_heading_added = True

                continue

            # 일반 단락
            para = doc.add_paragraph()
            self._add_inline_formatting(para, line)

    def _add_inline_formatting(self, paragraph, text: str):
        """인라인 마크다운 형식(볼드, 이탤릭) 처리

        Args:
            paragraph: Word 단락
            text: 텍스트
        """
        # 볼드+이탤릭: ***text*** 또는 ___text___
        # 볼드: **text** 또는 __text__
        # 이탤릭: *text* 또는 _text_

        pattern = r'(\*\*\*|___|__|\*\*|_|\*)(.*?)\1'
        last_pos = 0

        for match in re.finditer(pattern, text):
            # 일반 텍스트 추가
            if match.start() > last_pos:
                run = paragraph.add_run(text[last_pos:match.start()])
                run.font.size = Pt(11)
                run.font.name = 'Malgun Gothic'

            # 형식화된 텍스트 추가
            marker = match.group(1)
            content = match.group(2)

            run = paragraph.add_run(content)
            run.font.size = Pt(11)
            run.font.name = 'Malgun Gothic'

            if marker in ['***', '___']:
                run.bold = True
                run.italic = True
            elif marker in ['**', '__']:
                run.bold = True
            elif marker in ['*', '_']:
                run.italic = True

            last_pos = match.end()

        # 남은 텍스트 추가
        if last_pos < len(text):
            run = paragraph.add_run(text[last_pos:])
            run.font.size = Pt(11)
            run.font.name = 'Malgun Gothic'

    def generate_word_report(self, report_data: Dict[str, Any], output_path: str):
        """Word 보고서 생성

        Args:
            report_data: 보고서 데이터 (JSON)
            output_path: 출력 파일 경로
        """
        # Pandoc 사용 가능 여부 확인
        if PANDOC_AVAILABLE:
            try:
                self._generate_word_with_pandoc(report_data, output_path)
                return
            except Exception as e:
                print(f"⚠️ Pandoc 변환 실패, 기본 방식으로 전환: {e}")

        # 기본 방식 (python-docx)
        self._generate_word_basic(report_data, output_path)

    def _fix_table_format(self, text: str) -> str:
        """유니코드 박스 문자를 마크다운 테이블로 변환하고 구분선 추가"""
        # │ (유니코드 박스 문자)를 | (파이프)로 변환
        text = text.replace('│', '|')
        # ─ (유니코드 가로선)를 - (하이픈)로 변환
        text = text.replace('─', '-')

        # 테이블 구조 수정: 구분선이 없는 테이블에 구분선 추가
        lines = text.split('\n')
        fixed_lines = []
        i = 0

        while i < len(lines):
            line = lines[i]

            # 파이프로 시작하는 행 감지 (잠재적 테이블 행)
            if line.strip().startswith('|') and line.strip().endswith('|'):
                # 테이블 시작
                table_lines = [line]
                i += 1

                # 다음 행이 구분선인지 확인
                if i < len(lines):
                    next_line = lines[i].strip()

                    # 구분선이 아니면 (데이터 행이면) 구분선 추가
                    if next_line.startswith('|') and not re.match(r'^\s*\|[\s\-:]+\|\s*$', next_line):
                        # 첫 번째 행의 열 개수만큼 구분선 생성
                        num_cols = line.count('|') - 1
                        separator = '|' + '|'.join(['---' for _ in range(num_cols)]) + '|'
                        fixed_lines.append(line)
                        fixed_lines.append(separator)
                        continue

                fixed_lines.append(line)
            else:
                fixed_lines.append(line)
                i += 1

        return '\n'.join(fixed_lines)

    def _generate_word_with_pandoc(self, report_data: Dict[str, Any], output_path: str):
        """Pandoc을 사용한 Word 생성 (마크다운 완벽 지원)"""
        # 결과 수집
        results = report_data.get('results', [])
        markdown_content = []

        for result in results:
            if result.get('success'):
                answer = result.get('answer', 'N/A')
                # 테이블 형식 수정
                answer = self._fix_table_format(answer)
                markdown_content.append(answer)
                markdown_content.append('\n\n')  # 질문 사이 간격

        full_markdown = '\n'.join(markdown_content)

        # 임시 마크다운 파일 생성
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as tmp:
            tmp.write(full_markdown)
            tmp_path = tmp.name

        try:
            # Pandoc으로 변환
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            pypandoc.convert_file(
                tmp_path,
                'docx',
                outputfile=str(output_file),
                extra_args=['--reference-doc='] if False else []  # 필요시 템플릿 추가
            )

            print(f"📄 Word 문서 생성 완료 (Pandoc): {output_file}")

        finally:
            # 임시 파일 삭제
            Path(tmp_path).unlink(missing_ok=True)

    def _generate_word_basic(self, report_data: Dict[str, Any], output_path: str):
        """기본 방식으로 Word 생성 (python-docx)"""
        doc = Document()

        # 첫 번째 헤딩 플래그 초기화
        self.first_heading_added = False

        # 결과
        results = report_data.get('results', [])

        for i, result in enumerate(results, 1):
            if result.get('success'):
                # 답변만 표시 (질문과 검색된 문서 정보는 제외)
                answer = result.get('answer', 'N/A')
                # 테이블 형식 수정
                answer = self._fix_table_format(answer)

                # 마크다운 형식 처리
                self._add_formatted_text(doc, answer)

            else:
                # 오류 발생
                self._add_paragraph(doc, f"❌ 오류 발생: {result.get('error', 'Unknown error')}", bold=True)

            # 질문 사이 간격
            if i < len(results):
                doc.add_paragraph()
                doc.add_paragraph()

        # 저장
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_file)

        print(f"📄 Word 문서 생성 완료 (기본): {output_file}")

    def generate_pdf_report(self, report_data: Dict[str, Any], output_path: str):
        """PDF 보고서 생성 (Word를 먼저 만들고 PDF로 변환)

        Args:
            report_data: 보고서 데이터 (JSON)
            output_path: 출력 파일 경로
        """
        # 임시 Word 파일 생성
        temp_docx = output_path.replace('.pdf', '_temp.docx')
        self.generate_word_report(report_data, temp_docx)

        try:
            # Word를 PDF로 변환 (LibreOffice 사용)
            import subprocess

            output_file = Path(output_path)
            output_dir = output_file.parent

            # LibreOffice로 변환
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                temp_docx
            ]

            result = subprocess.run(cmd, capture_output=True, text=True)

            if result.returncode == 0:
                # 생성된 PDF 파일을 원하는 이름으로 변경
                generated_pdf = temp_docx.replace('.docx', '.pdf')
                if Path(generated_pdf).exists() and generated_pdf != output_path:
                    Path(generated_pdf).rename(output_path)

                print(f"📄 PDF 문서 생성 완료: {output_path}")
            else:
                print(f"⚠️ PDF 변환 실패: {result.stderr}")
                print(f"💡 Word 파일을 사용하세요: {temp_docx}")

        except Exception as e:
            print(f"⚠️ PDF 생성 실패: {e}")
            print(f"💡 Word 파일을 사용하세요: {temp_docx}")

        finally:
            # 임시 Word 파일 삭제 (옵션)
            # Path(temp_docx).unlink(missing_ok=True)
            pass


def main():
    """테스트용 메인 함수"""
    import json
    import argparse

    parser = argparse.ArgumentParser(description="보고서 문서 생성기")
    parser.add_argument("--json", type=str, required=True, help="입력 JSON 파일")
    parser.add_argument("--output", type=str, required=True, help="출력 파일 (.docx 또는 .pdf)")

    args = parser.parse_args()

    # JSON 로드
    with open(args.json, 'r', encoding='utf-8') as f:
        report_data = json.load(f)

    # 문서 생성
    generator = DocumentGenerator()

    if args.output.endswith('.pdf'):
        generator.generate_pdf_report(report_data, args.output)
    elif args.output.endswith('.docx'):
        generator.generate_word_report(report_data, args.output)
    else:
        print("❌ 지원되지 않는 파일 형식입니다. .docx 또는 .pdf를 사용하세요.")


if __name__ == "__main__":
    main()
